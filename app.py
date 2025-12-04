import streamlit as st
from docx import Document
from datetime import date, datetime
from io import BytesIO
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import re

st.set_page_config(page_title="Advers Bildirim v22", page_icon="🇹🇷", layout="centered")

# --- AYARLAR ---
GONDEREN_EMAIL = "mersinfarmakoloji@gmail.com"
ALICI_EMAIL = "mersinfarmakoloji@gmail.com"

st.title("🇹🇷 T.C. Sağlık Bakanlığı - TÜFAM Bildirimi")

st.warning("⚠️ Gönderim için; Hasta Adı, En az bir İlaç, En az bir Reaksiyon, Bildirimi Yapan Doktorun Adı ve Telefon numarası ZORUNLUDUR.")

# --- YARDIMCI FONKSİYONLAR ---
def tr_to_en_filename(text):
    if not text: return "Rapor"
    mapping = {
        'ç': 'c', 'Ç': 'C', 'ğ': 'g', 'Ğ': 'G', 'ı': 'i', 'I': 'I', 'İ': 'I',
        'ö': 'o', 'Ö': 'O', 'ş': 's', 'Ş': 'S', 'ü': 'u', 'Ü': 'U'
    }
    for k, v in mapping.items():
        text = text.replace(k, v)
    return text

def tarih_kontrol_ve_duzelt(girdi):
    if not girdi: return None
    girdi = girdi.strip().lower()
    if girdi in ["bugün", "bugun", "today"]:
        return date.today().strftime("%d.%m.%Y")
    if girdi.isdigit() and len(girdi) == 8:
        girdi = f"{girdi[:2]}.{girdi[2:4]}.{girdi[4:]}"
    girdi = girdi.replace("/", ".").replace("-", ".")
    try:
        datetime.strptime(girdi, "%d.%m.%Y")
        return girdi 
    except ValueError:
        return "HATA" 

def kutu_yap(secim, hedef):
    if secim is None: return "[ ]"
    return "[X]" if secim == hedef else "[ ]"

def soru_cevapla(cevap):
    if cevap == "Evet": return "[X] Evet  [ ] Hayır  [ ] Bilinmiyor"
    if cevap == "Hayır": return "[ ] Evet  [X] Hayır  [ ] Bilinmiyor"
    if cevap == "Bilinmiyor": return "[ ] Evet  [ ] Hayır  [X] Bilinmiyor"
    return "[ ] Evet  [ ] Hayır  [ ] Bilinmiyor"

def TR_upper(text):
    if text: return text.replace("i", "İ").upper()
    return ""

def TR_lower(text):
    if text: return text.replace("I", "ı").replace("İ", "i").lower()
    return ""

# --- A. HASTAYA AİT BİLGİLER ---
st.header("A. HASTAYA AİT BİLGİLER")
c1, c2 = st.columns(2)
with c1:
    ad_soyad = st.text_input("1. Hasta Ad Soyad (Baş Harfler)", placeholder="Örn: A.Y.")
    
    dogum_tarihi_raw = st.text_input("2. Doğum Tarihi", placeholder="GünAyYıl (Örn: 01011980)")
    dogum_tarihi = tarih_kontrol_ve_duzelt(dogum_tarihi_raw)
    
    yas_str = ""
    if dogum_tarihi == "HATA":
        st.error("❌ Geçersiz Tarih!")
        dogum_tarihi = "" 
    elif dogum_tarihi:
        try:
            dt_obj = datetime.strptime(dogum_tarihi, "%d.%m.%Y")
            bugun = date.today()
            yas_hesap = bugun.year - dt_obj.year - ((bugun.month, bugun.day) < (dt_obj.month, dt_obj.day))
            st.success(f"📅 Algılandı: {dogum_tarihi} (Yaş: {yas_hesap})")
            yas_str = str(yas_hesap)
        except: pass

with c2:
    cinsiyet = st.radio("3. Cinsiyet", ["Kadın", "Erkek"], horizontal=True, index=None)
    boy = st.text_input("4. Boy (cm)", placeholder="170")
    kilo = st.text_input("5. Ağırlık (kg)", placeholder="70")

st.markdown("---")
st.subheader("⚠️ Ciddiyet Durumu")

ciddiyet_durumu = st.radio("Vaka Ciddi mi?", ["Ciddi Değil", "Ciddi"], horizontal=True, index=None)

k_olum_val, k_hayat_val, k_hastane_val, k_sakatlik_val, k_anomali_val, k_tibbi_val = False, False, False, False, False, False
olum_tarihi_str, olum_nedeni, otopsi = "", "", "[ ] Evet  [ ] Hayır"

if ciddiyet_durumu == "Ciddi":
    with st.container():
        cols_cid = st.columns(2)
        with cols_cid[0]:
            k_olum_val = st.checkbox("💀 Ölüm")
            k_hayat_val = st.checkbox("❤️ Hayatı Tehdit Edici")
            k_hastane_val = st.checkbox("🏥 Hastaneye Yatış/Uzama")
        with cols_cid[1]:
            k_sakatlik_val = st.checkbox("♿ Kalıcı Sakatlık")
            k_anomali_val = st.checkbox("👶 Konjenital Anomali")
            k_tibbi_val = st.checkbox("⚕️ Tıbbi Olarak Önemli")

    if k_olum_val:
        col_o1, col_o2 = st.columns(2)
        with col_o1:
            ot_raw = st.text_input("Ölüm Tarihi", placeholder="GünAyYıl")
            olum_tarihi_str = tarih_kontrol_ve_duzelt(ot_raw)
            if olum_tarihi_str == "HATA": 
                st.error("Geçersiz Tarih")
                olum_tarihi_str = ""
            
            oto = st.radio("Otopsi Yapıldı mı?", ["Evet", "Hayır"], horizontal=True, index=None)
            
            if oto == "Evet": otopsi = "[X] Evet  [ ] Hayır"
            elif oto == "Hayır": otopsi = "[ ] Evet  [X] Hayır"
            else: otopsi = "[ ] Evet  [ ] Hayır"

        with col_o2:
            olum_nedeni = st.text_input("Ölüm Nedeni")

# --- B. REAKSİYONLAR ---
st.header("B. ADVERS REAKSİYONLAR")
reaksiyonlar = []

# İlk reaksiyonun tarihlerini hafızada tutmak için değişkenler
ilk_r_bas = ""
ilk_r_bit = ""

for i in range(1, 6):
    with st.expander(f"Reaksiyon {i}", expanded=(i==1)):
        col_r1, col_r2, col_r3 = st.columns([3, 1, 1])
        with col_r1: r_tanim = st.text_input(f"Tanım", key=f"rt{i}")
        
        # --- BAŞLANGIÇ TARİHİ ALANI ---
        with col_r2: 
            r_bas = ""
            use_first_bas = False
            
            # 2. ve sonraki satırlar için 'Kopyala' kutucuğu
            if i > 1:
                use_first_bas = st.checkbox("1. ile aynı", key=f"r_bas_copy_{i}")
            
            if use_first_bas:
                # Eğer kutu işaretliyse, ilk değeri al ve ekrana bilgi yaz (Input gizlenir)
                r_bas = ilk_r_bas
                st.caption(f"🗓️ {ilk_r_bas}")
            else:
                # İşaretli değilse veya 1. satırsa normal giriş
                rb_raw = st.text_input(f"Başlangıç", key=f"rb{i}", placeholder="GünAyYıl")
                r_bas = tarih_kontrol_ve_duzelt(rb_raw)
                if r_bas == "HATA": st.error("Tarih Hatalı"); r_bas=""
            
            # Eğer 1. satırsak, bu değeri hafızaya at
            if i == 1: ilk_r_bas = r_bas

        # --- BİTİŞ TARİHİ ALANI ---
        with col_r3: 
            r_devam = st.checkbox("Devam Ediyor", key=f"rd{i}")
            if r_devam:
                r_bit = "DEVAM EDİYOR"
                # Devam ediyorsa hafızaya da öyle kaydet
                if i == 1: ilk_r_bit = "DEVAM EDİYOR"
            else:
                r_bit = ""
                use_first_bit = False
                
                # 2. ve sonrası için 'Kopyala' kutucuğu
                if i > 1:
                    use_first_bit = st.checkbox("1. ile aynı", key=f"r_bit_copy_{i}")
                
                if use_first_bit:
                    r_bit = ilk_r_bit
                    st.caption(f"🗓️ {ilk_r_bit}")
                else:
                    rbit_raw = st.text_input(f"Bitiş", key=f"rbit{i}", placeholder="GünAyYıl")
                    r_bit = tarih_kontrol_ve_duzelt(rbit_raw)
                    if r_bit == "HATA": st.error("Tarih Hatalı"); r_bit=""
                
                if i == 1: ilk_r_bit = r_bit

        if r_tanim: 
            reaksiyonlar.append({"tanim": r_tanim, "bas": r_bas, "bit": r_bit, "devam": r_devam})

st.subheader("Sonuç Durumu")
sonuc_secim = st.radio("Sonuç", ["İyileşti/Düzeldi", "İyileşiyor", "Sekel Bıraktı", "Devam Ediyor", "Ölümle Sonuçlandı", "Bilinmiyor"], horizontal=True, index=None)

lab_bulgu = st.text_area("3. Laboratuvar Bulguları (Tarihleriyle birlikte)", height=68)
st.info("ℹ️ **Tıbbi Öykü:** Allerji, gebelik, sigara/alkol, kronik hastalıklar vb.")
tibbi_oyku = st.text_area("4. Tıbbi Öykü / Eş Zamanlı Hastalıklar", height=68)

# --- C. İLAÇLAR ---
st.header("C. ŞÜPHELENİLEN İLAÇLAR")
ilaclar = []

# İlk ilacın tarihlerini hafızada tutmak için değişkenler
ilk_i_bas = ""
ilk_i_bit = ""

for i in range(1, 6):
    with st.expander(f"💊 İlaç {i}", expanded=(i==1)):
        c_i1, c_i2, c_i3 = st.columns([2, 1, 1])
        with c_i1: 
            i_adi = st.text_input(f"İlaç Adı", key=f"ia{i}", help="Biliniyorsa TİCARİ ismini yazınız.")
        with c_i2: 
            i_yol_secim = st.selectbox(f"Veriliş Yolu", ["Oral", "IV", "IM", "SC", "Topikal", "Diğer"], key=f"iy{i}")
            if i_yol_secim == "Diğer":
                i_yol = st.text_input(f"👉 Yolu Yazınız ({i})", key=f"iy_txt{i}")
            else:
                i_yol = i_yol_secim
        with c_i3: 
            i_doz = st.text_input(f"Günlük Doz", placeholder="Örn: 500 mg", key=f"id{i}")
        
        c_i4, c_i5, c_i6 = st.columns([2, 1, 1])
        with c_i4: i_end = st.text_input(f"Endikasyon", key=f"ie{i}")
        
        # --- İLAÇ BAŞLAMA TARİHİ ---
        with c_i5: 
            i_bas = ""
            use_first_ibase = False
            
            if i > 1:
                use_first_ibase = st.checkbox("1. ile aynı", key=f"i_bas_copy_{i}")
            
            if use_first_ibase:
                i_bas = ilk_i_bas
                st.caption(f"🗓️ {ilk_i_bas}")
            else:
                ib_raw = st.text_input(f"Başlama", key=f"ib{i}", placeholder="GünAyYıl")
                i_bas = tarih_kontrol_ve_duzelt(ib_raw)
                if i_bas == "HATA": st.error("Geçersiz Tarih"); i_bas=""
            
            if i == 1: ilk_i_bas = i_bas

        # --- İLAÇ KESİLME TARİHİ ---
        with c_i6: 
            i_devam = st.checkbox("Kullanım Devam Ediyor", key=f"idvm{i}")
            if i_devam:
                i_bit = "DEVAM EDİYOR"
                if i == 1: ilk_i_bit = "DEVAM EDİYOR"
            else:
                i_bit = ""
                use_first_ibit = False
                
                if i > 1:
                    use_first_ibit = st.checkbox("1. ile aynı", key=f"i_bit_copy_{i}")
                
                if use_first_ibit:
                    i_bit = ilk_i_bit
                    st.caption(f"🗓️ {ilk_i_bit}")
                else:
                    ibit_raw = st.text_input(f"Kesilme", key=f"ibit{i}", placeholder="GünAyYıl")
                    i_bit = tarih_kontrol_ve_duzelt(ibit_raw)
                    if i_bit == "HATA": st.error("Geçersiz Tarih"); i_bit=""
                
                if i == 1: ilk_i_bit = i_bit

        st.markdown(f":blue[**⬇️ {i}. İlaç Değerlendirme Soruları:**]")
        q7 = st.radio("7. İlaç Kesildi mi?", ["Evet", "Hayır", "Bilinmiyor"], key=f"q7_{i}", horizontal=True, index=None)
        q8 = st.radio("8. Reaksiyon azaldı mı?", ["Evet", "Hayır", "Bilinmiyor"], key=f"q8_{i}", horizontal=True, index=None)
        q9 = st.radio("9. Yeniden verildi mi?", ["Evet", "Hayır", "Bilinmiyor"], key=f"q9_{i}", horizontal=True, index=None)
        q10 = st.radio("10. Tekrarladı mı?", ["Evet", "Hayır", "Bilinmiyor"], key=f"q10_{i}", horizontal=True, index=None)

        if i_adi: 
            ilaclar.append({
                "ad": i_adi, "yol": i_yol, "doz": i_doz, "bas": i_bas, "bit": i_bit, "end": i_end,
                "s7": soru_cevapla(q7), "s8": soru_cevapla(q8), "s9": soru_cevapla(q9), "s10": soru_cevapla(q10),
                "devam": i_devam
            })

st.info("ℹ️ Eş Zamanlı ilaçları virgül ile ayırarak yazınız.")
es_zamanli = st.text_area("11. Eş Zamanlı İlaçlar", height=68)
diger_gozlem = st.text_area("12. Diğer Gözlemler (Kalite sorunu vb.)", height=68)
tedavi = st.text_area("13. Advers Reaksiyonun Tedavisi", height=68)

# --- D. BİLDİREN ---
st.header("D. BİLDİRİM YAPAN KİŞİ")
c_d1, c_d2 = st.columns(2)
with c_d1:
    b_ad = st.text_input("1. Adı Soyadı (Bildirimi Yapan)")
    b_tel = st.text_input("3. Tel No")
    b_faks = st.text_input("5. Faks")
with c_d2:
    b_meslek = st.radio("2. Meslek", ["Doktor", "Eczacı", "Hemşire", "Diğer"], horizontal=True, index=None)
    b_adres = st.text_area("4. Adres ve Bölüm", value="Mersin Üniversitesi Tıp Fakültesi", height=100)
    b_email = st.text_input("6. E-posta")

st.markdown("---")
col_r1, col_r2 = st.columns(2)
with col_r1:
    rapor_firma = st.radio("8. Rapor firmaya bildirildi mi?", ["Bilinmiyor", "Evet", "Hayır"], horizontal=True, index=None)
with col_r2:
    rapor_tipi = st.radio("10. Rapor Tipi", ["İlk", "Takip"], horizontal=True, index=None)

rt_raw = st.text_input("9. Rapor Tarihi", value=date.today().strftime("%d.%m.%Y"))
rapor_tarihi = tarih_kontrol_ve_duzelt(rt_raw)

st.markdown("---")
submitted = st.button("📤 BİLDİRİMİ GÖNDER", type="primary", use_container_width=True)

# --- KAYIT VE MAİL ---
if submitted:
    eksik_alanlar = []

    if not ad_soyad: eksik_alanlar.append("Hasta Adı Soyadı")
    if not ilaclar: eksik_alanlar.append("En az bir İlaç Adı")
    if not reaksiyonlar: eksik_alanlar.append("En az bir Reaksiyon Tanımı")
    if not b_ad: eksik_alanlar.append("Bildirimi Yapan Kişi Adı")
    if not b_tel: eksik_alanlar.append("Bildirimi Yapan Telefon No")
    if not b_meslek: eksik_alanlar.append("Meslek Seçimi")

    if len(eksik_alanlar) > 0:
        st.error("⚠️ GÖNDERİM BAŞARISIZ! Lütfen aşağıdaki eksik alanları doldurunuz:")
        for eksik in eksik_alanlar:
            st.warning(f"❌ {eksik} eksik.")
    
    else:
        try:
            with st.spinner("Rapor oluşturuluyor ve mail gönderiliyor..."):
                doc = Document("Advers reaksiyon bildirim formu.docx")
                
                # --- VERİ HAZIRLIĞI ---
                r_list = [{"tanim":"", "bas":"", "bit":""} for _ in range(5)]
                for idx, r in enumerate(reaksiyonlar):
                    bitis_str = "DEVAM EDİYOR" if r["devam"] else r["bit"]
                    if idx < 5:
                        r_list[idx] = {"tanim": TR_upper(r["tanim"]), "bas": r["bas"], "bit": bitis_str}

                i_list = [{"ad":"", "yol":"", "doz":"", "bas":"", "bit":"", "end":"", "s7":"", "s8":"", "s9":"", "s10":""} for _ in range(5)]
                for idx, ilac in enumerate(ilaclar):
                    bitis_str = "DEVAM EDİYOR" if ilac["devam"] else ilac["bit"]
                    if idx < 5:
                        i_list[idx] = {
                            "ad": TR_upper(ilac["ad"]), "yol": TR_upper(ilac["yol"]), "doz": TR_lower(ilac["doz"]), 
                            "end": TR_upper(ilac["end"]), "bas": ilac["bas"], "bit": bitis_str,
                            "s7": ilac["s7"], "s8": ilac["s8"], "s9": ilac["s9"], "s10": ilac["s10"]
                        }

                def radio_kutu(secim, hedef): 
                    if secim is None: return "[ ]"
                    return "[X]" if secim == hedef else "[ ]"

                rf_str = "[ ] Evet [ ] Hayır [ ] Bilinmiyor" if rapor_firma is None else f"{radio_kutu(rapor_firma, 'Evet')} Evet  {radio_kutu(rapor_firma, 'Hayır')} Hayır  {radio_kutu(rapor_firma, 'Bilinmiyor')} Bilinmiyor"
                rt_str = "[ ] İlk [ ] Takip" if rapor_tipi is None else f"{radio_kutu(rapor_tipi, 'İlk')} İlk  {radio_kutu(rapor_tipi, 'Takip')} Takip"

                veriler = {
                    "{{hasta_adi_soyadi_basharfleri}}": TR_upper(ad_soyad), 
                    "{{dogum_tarihi}}": dogum_tarihi, "{{yas}}": yas_str, 
                    "{{cinsiyet}}": cinsiyet if cinsiyet else "",
                    "{{boy}}": boy, "{{kilo}}": kilo,
                    "{{cid_yok}}": "[X]" if ciddiyet_durumu == "Ciddi Değil" else "[ ]", "{{cid_var}}": "[X]" if ciddiyet_durumu == "Ciddi" else "[ ]",
                    "{{k_olum}}": "[X]" if k_olum_val else "[ ]", "{{k_hayat}}": "[X]" if k_hayat_val else "[ ]",
                    "{{k_hastane}}": "[X]" if k_hastane_val else "[ ]", "{{k_sakatlik}}": "[X]" if k_sakatlik_val else "[ ]",
                    "{{k_anomali}}": "[X]" if k_anomali_val else "[ ]", "{{k_tibbi}}": "[X]" if k_tibbi_val else "[ ]",
                    "{{olum_tarih}}": olum_tarihi_str, "{{olum_neden}}": TR_upper(olum_nedeni), "{{otopsi}}": otopsi,
                    "{{reaksiyon_1}}": r_list[0]["tanim"], "{{bas_1}}": r_list[0]["bas"], "{{bit_1}}": r_list[0]["bit"],
                    "{{reaksiyon_2}}": r_list[1]["tanim"], "{{bas_2}}": r_list[1]["bas"], "{{bit_2}}": r_list[1]["bit"],
                    "{{reaksiyon_3}}": r_list[2]["tanim"], "{{bas_3}}": r_list[2]["bas"], "{{bit_3}}": r_list[2]["bit"],
                    "{{reaksiyon_4}}": r_list[3]["tanim"], "{{bas_4}}": r_list[3]["bas"], "{{bit_4}}": r_list[3]["bit"],
                    "{{reaksiyon_5}}": r_list[4]["tanim"], "{{bas_5}}": r_list[4]["bas"], "{{bit_5}}": r_list[4]["bit"],
                    "{{s_iyilesti}}": kutu_yap(sonuc_secim, "İyileşti/Düzeldi"), "{{s_iyilesiyor}}": kutu_yap(sonuc_secim, "İyileşiyor"), "{{s_sekel}}": kutu_yap(sonuc_secim, "Sekel Bıraktı"),
                    "{{s_devam}}": kutu_yap(sonuc_secim, "Devam Ediyor"), "{{s_olum}}": kutu_yap(sonuc_secim, "Ölümle Sonuçlandı"), "{{s_bilinmiyor}}": kutu_yap(sonuc_secim, "Bilinmiyor"),
                    "{{lab}}": TR_upper(lab_bulgu), "{{oyku}}": TR_upper(tibbi_oyku), "{{tedavi}}": TR_upper(tedavi), "{{diger_gozlem}}": TR_upper(diger_gozlem),
                    "{{ilac_1}}": i_list[0]["ad"], "{{yol_1}}": i_list[0]["yol"], "{{doz_1}}": i_list[0]["doz"], "{{ilac_bas_1}}": i_list[0]["bas"], "{{ilac_bit_1}}": i_list[0]["bit"], "{{end_1}}": i_list[0]["end"], "{{s7_1}}": i_list[0]["s7"], "{{s8_1}}": i_list[0]["s8"], "{{s9_1}}": i_list[0]["s9"], "{{s10_1}}": i_list[0]["s10"],
                    "{{ilac_2}}": i_list[1]["ad"], "{{yol_2}}": i_list[1]["yol"], "{{doz_2}}": i_list[1]["doz"], "{{ilac_bas_2}}": i_list[1]["bas"], "{{ilac_bit_2}}": i_list[1]["bit"], "{{end_2}}": i_list[1]["end"], "{{s7_2}}": i_list[1]["s7"], "{{s8_2}}": i_list[1]["s8"], "{{s9_2}}": i_list[1]["s9"], "{{s10_2}}": i_list[1]["s10"],
                    "{{ilac_3}}": i_list[2]["ad"], "{{yol_3}}": i_list[2]["yol"], "{{doz_3}}": i_list[2]["doz"], "{{ilac_bas_3}}": i_list[2]["bas"], "{{ilac_bit_3}}": i_list[2]["bit"], "{{end_3}}": i_list[2]["end"], "{{s7_3}}": i_list[2]["s7"], "{{s8_3}}": i_list[2]["s8"], "{{s9_3}}": i_list[2]["s9"], "{{s10_3}}": i_list[2]["s10"],
                    "{{ilac_4}}": i_list[3]["ad"], "{{yol_4}}": i_list[3]["yol"], "{{doz_4}}": i_list[3]["doz"], "{{ilac_bas_4}}": i_list[3]["bas"], "{{ilac_bit_4}}": i_list[3]["bit"], "{{end_4}}": i_list[3]["end"], "{{s7_4}}": i_list[3]["s7"], "{{s8_4}}": i_list[3]["s8"], "{{s9_4}}": i_list[3]["s9"], "{{s10_4}}": i_list[3]["s10"],
                    "{{ilac_5}}": i_list[4]["ad"], "{{yol_5}}": i_list[4]["yol"], "{{doz_5}}": i_list[4]["doz"], "{{ilac_bas_5}}": i_list[4]["bas"], "{{ilac_bit_5}}": i_list[4]["bit"], "{{end_5}}": i_list[4]["end"], "{{s7_5}}": i_list[4]["s7"], "{{s8_5}}": i_list[4]["s8"], "{{s9_5}}": i_list[4]["s9"], "{{s10_5}}": i_list[4]["s10"],
                    "{{bildiren_ad}}": TR_upper(b_ad), "{{bildiren_meslek}}": b_meslek if b_meslek else "", 
                    "{{bildiren_tel}}": b_tel, 
                    "{{bildiren_adres}}": TR_upper(b_adres), "{{bildiren_faks}}": b_faks, "{{bildiren_email}}": b_email,
                    "{{rapor_tarihi}}": rapor_tarihi,
                    "{{rapor_firma}}": rf_str, "{{rapor_tipi}}": rt_str,
                    "{{es_zamanli}}": TR_upper(es_zamanli)
                }

                def replace_text_preserving_style(doc, data):
                    for p in doc.paragraphs:
                        if "{{" in p.text: 
                            for key, value in data.items():
                                if key in p.text: p.text = p.text.replace(key, str(value))
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    if "{{" in p.text:
                                        for key, value in data.items():
                                            if key in p.text: p.text = p.text.replace(key, str(value))
                    regex = re.compile(r"\{\{.*?\}\}") 
                    for p in doc.paragraphs:
                        if "{{" in p.text: p.text = regex.sub("", p.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    if "{{" in p.text: p.text = regex.sub("", p.text)

                replace_text_preserving_style(doc, veriler)
                bio = BytesIO()
                doc.save(bio)
                
                # --- MAİL GÖNDERME ---
                try:
                    GMAIL_SIFRE = st.secrets["GMAIL_PASS"] 
                    msg = MIMEMultipart()
                    msg['From'] = GONDEREN_EMAIL
                    msg['To'] = ALICI_EMAIL
                    
                    clean_filename = f"Advers_{tr_to_en_filename(ad_soyad)}.docx"
                    
                    msg['Subject'] = f"Advers Raporu - {TR_upper(ad_soyad)}"
                    body = f"Sayın Yetkili,\n\n{TR_upper(ad_soyad)} hastasına ait rapor ektedir."
                    msg.attach(MIMEText(body, 'plain'))
                    part = MIMEBase('application', "octet-stream")
                    part.set_payload(bio.getvalue())
                    encoders.encode_base64(part)
                    part.add_header('Content-Disposition', f'attachment; filename="{clean_filename}"')
                    msg.attach(part)
                    server = smtplib.SMTP('smtp.gmail.com', 587)
                    server.starttls()
                    server.login(GONDEREN_EMAIL, GMAIL_SIFRE)
                    server.sendmail(GONDEREN_EMAIL, ALICI_EMAIL, msg.as_string())
                    server.quit()
                    st.success(f"✅ Rapor başarıyla {ALICI_EMAIL} adresine gönderildi!")
                except Exception as mail_err:
                    st.warning(f"⚠️ Mail gönderilemedi. (Sebep: {mail_err})")
                
                st.download_button(label="📥 RAPORU İNDİR", data=bio.getvalue(), file_name=f"Advers_{ad_soyad}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
        except Exception as e:
            st.error(f"Hata: {e}")
